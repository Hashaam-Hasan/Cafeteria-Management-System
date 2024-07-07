from django.core.management.base import BaseCommand
from django.db import connection
from docx import Document
from docx.shared import Pt

class Command(BaseCommand):
    help = 'Generate a report of all indexes for all tables in the database'

    def handle(self, *args, **options):
        # Create a new Document
        doc = Document()
        
        # Title page
        doc.add_heading('Database Index Report', 0)
        doc.add_paragraph('Generated using Django Management Command')
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('Table of Contents', level=1)
        toc = doc.add_paragraph('')
        toc_run = toc.add_run()
        toc_run.add_break()
        
        # Collect index information
        index_info = {}
        with connection.cursor() as cursor:
            tables = connection.introspection.table_names(cursor)
            for table_name in tables:
                constraints = connection.introspection.get_constraints(cursor, table_name)
                index_info[table_name] = []
                for constraint_name, constraint_details in constraints.items():
                    if constraint_details['index']:
                        index_info[table_name].append((constraint_name, constraint_details))
        
        # Add indexes information
        for table_name, indexes in index_info.items():
            doc.add_heading(f'Indexes for table {table_name}', level=2)
            toc_run.add_text(f'Indexes for table {table_name}\n')
            
            for index_name, index_details in indexes:
                doc.add_heading(f'- {index_name}', level=3)
                for key, value in index_details.items():
                    doc.add_paragraph(f'{key}: {value}', style='ListBullet')
        
        # Save the document
        output_file_path = 'indexes_report.docx'
        doc.save(output_file_path)
        self.stdout.write(f"Indexes report has been written to {output_file_path}")
